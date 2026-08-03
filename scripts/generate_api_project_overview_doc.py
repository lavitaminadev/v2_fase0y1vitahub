from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "current" / "VITAHUB_APIs_Arquitectura_y_Flujo.docx"

TITLE = "VITAHUB: APIs, Arquitectura del Proyecto y Flujo Operativo"
SUBTITLE = (
    "Resumen tecnico breve para presentacion de Fase 0 y Fase 1, "
    "enfocado en endpoints, organizacion de archivos y conexion CRM -> reservas -> Meta."
)

TEXT = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
ACCENT = RGBColor(0x00, 0x00, 0x00)
LIGHT = "F2F4F7"


API_ROWS = [
    (
        "Autenticacion",
        "/api/auth",
        "Login, refresh, recuperacion y control de sesion.",
        "core/auth/auth.controller.ts",
    ),
    (
        "Clientes",
        "/api/clients",
        "Gestion de cuentas cliente administradas por La Vitamina.",
        "modules/clients/clients.controller.ts",
    ),
    (
        "CRM leads",
        "/api/crm/leads",
        "Ingreso, consulta, actualizacion y conversion de leads.",
        "modules/crm/leads/lead.controller.ts",
    ),
    (
        "CRM contactos",
        "/api/crm/contacts",
        "Contactos consolidados del CRM por cliente.",
        "modules/crm/contacts/contacts.controller.ts",
    ),
    (
        "CRM oportunidades",
        "/api/crm/opportunities",
        "Seguimiento comercial y operativo por oportunidad.",
        "modules/crm/opportunities/opportunities.controller.ts",
    ),
    (
        "CRM interacciones",
        "/api/crm/interactions",
        "Bitacora de llamadas, notas, mensajes y seguimiento.",
        "modules/crm/interactions/interactions.controller.ts",
    ),
    (
        "Reservas internas",
        "/api/reservations",
        "Configuracion, gestion, estados y administracion de reservas.",
        "modules/reservations/reservations.controller.ts",
    ),
    (
        "Reservas publicas",
        "/api/public/reservations",
        "Entrada publica desde formulario o landing de reservas.",
        "modules/reservations/public-reservations.controller.ts",
    ),
    (
        "Meta e integraciones",
        "/api/integrations/meta",
        "Conexion, sincronizacion y eventos asociados a Meta.",
        "modules/integrations/meta/meta-pixel.controller.ts",
    ),
    (
        "Reportes",
        "/api/reporting",
        "Indicadores, resultados y lectura consolidada por cliente.",
        "modules/reports/reports.controller.ts",
    ),
]

PROJECT_ROWS = [
    (
        "apps/api",
        "Backend principal NestJS",
        "Controllers, services/use-cases, entities, DTOs, auth, permisos, integraciones y reglas de negocio.",
    ),
    (
        "apps/web",
        "Frontend principal React + Vite",
        "Vistas por feature: CRM, reservations, integrations, dashboard, portal cliente, settings y usuarios.",
    ),
    (
        "database",
        "Base de datos y soporte relacional",
        "Schema, seeds, diagramas ER y documentacion tecnica de tablas y relaciones.",
    ),
    (
        "docs",
        "Documentacion funcional y tecnica",
        "Fases, arquitectura, API, CRM, reservas, integraciones y materiales de presentacion.",
    ),
    (
        "scripts",
        "Automatizaciones y utilidades",
        "Seeds, generadores documentales, ayudas de entorno y tareas de soporte.",
    ),
    (
        "packages",
        "Codigo compartido",
        "Espacio para utilidades comunes y piezas reutilizables entre apps.",
    ),
    (
        "infrastructure",
        "Soporte de despliegue y entorno",
        "Archivos de apoyo tecnico para hosting, operacion y configuracion.",
    ),
]

MODULE_ROWS = [
    (
        "CRM Leads",
        "lead.controller.ts, create/update/list DTOs, use-cases, lead-intake.service.ts",
        "Captura y procesa leads de formularios, campañas o carga manual.",
    ),
    (
        "CRM Pipeline",
        "contacts, opportunities, interactions",
        "Organiza seguimiento del lead ya identificado y lo vincula a cliente y operacion.",
    ),
    (
        "Reservas",
        "reservations.controller.ts, public-reservations.controller.ts, reservations.service.ts, entities",
        "Convierte interes o agenda en reserva trazable con formulario, estado y eventos.",
    ),
    (
        "Portal Cliente",
        "apps/web/src/features/client-portal",
        "Expone al cliente su vista acotada de avances, reportes, reuniones y elementos propios.",
    ),
    (
        "Integraciones",
        "integrations/meta, integrations/google",
        "Conecta el ecosistema de adquisicion y reporte con plataformas externas.",
    ),
]

FRONTEND_ROWS = [
    (
        "CRM",
        "apps/web/src/features/crm",
        "LeadsPage.tsx, CrmRecordsPage.tsx, LeadDetailDrawer.tsx",
    ),
    (
        "Reservas",
        "apps/web/src/features/reservations",
        "ReservationsPage.tsx, ReservationBuilderPage.tsx, PublicReservationPage.tsx",
    ),
    (
        "Integraciones",
        "apps/web/src/features/integrations",
        "IntegrationsPage.tsx, MetaConnectCard.tsx, OAuthCallbackPage.tsx",
    ),
    (
        "Portal cliente",
        "apps/web/src/features/client-portal",
        "ClientDashboard.tsx, ClientReports.tsx, ClientMeetings.tsx, ClientApprovals.tsx",
    ),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_font(run, name="Calibri", size=11, bold=False, color=TEXT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def style_paragraph(paragraph, after=6, before=0, line=1.15, align=WD_ALIGN_PARAGRAPH.LEFT):
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(after)
    fmt.space_before = Pt(before)
    fmt.line_spacing = line
    paragraph.alignment = align


def add_title(doc):
    p = doc.add_paragraph()
    style_paragraph(p, after=3)
    r = p.add_run(TITLE)
    set_font(r, name="Arial", size=22, bold=False, color=ACCENT)

    p2 = doc.add_paragraph()
    style_paragraph(p2, after=14)
    r2 = p2.add_run(SUBTITLE)
    set_font(r2, name="Arial", size=10, color=MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style_paragraph(p, before=14 if level == 1 else 10, after=6)
    r = p.add_run(text)
    set_font(r, name="Arial", size=15 if level == 1 else 12, bold=False, color=ACCENT if level == 1 else TEXT)


def add_body(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, after=6)
    r = p.add_run(text)
    set_font(r, name="Arial", size=11, color=TEXT)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        style_paragraph(p, after=4)
        r = p.add_run(item)
        set_font(r, name="Arial", size=11, color=TEXT)


def fill_table(table, headers, rows, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, value in enumerate(headers):
        header_cells[idx].width = widths[idx]
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(header_cells[idx], LIGHT)
        p = header_cells[idx].paragraphs[0]
        style_paragraph(p, after=0)
        r = p.add_run(value)
        set_font(r, name="Arial", size=10, bold=True)

    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].width = widths[idx]
            row[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = row[idx].paragraphs[0]
            style_paragraph(p, after=0)
            r = p.add_run(value)
            set_font(r, name="Arial", size=10)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.start_type = WD_SECTION_START.NEW_PAGE

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)

    add_title(doc)

    add_heading(doc, "1. Lectura Ejecutiva")
    add_body(
        doc,
        "VITAHUB esta organizado como una plataforma modular donde el backend expone APIs por dominio "
        "y el frontend consume esas capacidades por vistas funcionales. La base tecnica mantiene "
        "aislamiento por organization_id, interpretado funcionalmente como La Vitamina operando "
        "multiples clientes dentro del mismo ecosistema."
    )
    add_bullets(
        doc,
        [
            "Prefijo global de API: /api",
            "Backend principal: NestJS por modulos",
            "Frontend principal: React + Vite por features",
            "Flujo objetivo de Fase 1: lead -> seguimiento CRM -> reserva -> asistencia -> Meta/reporting",
        ],
    )

    add_heading(doc, "2. Mapa de APIs y Endpoints")
    add_body(
        doc,
        "Los endpoints estan agrupados por dominio. Cada bloque suele tener controller para HTTP, "
        "DTOs para entradas, entities para persistencia y services o use-cases para la logica."
    )
    table = doc.add_table(rows=1, cols=4)
    fill_table(
        table,
        ["Dominio", "Endpoint base", "Uso principal", "Archivo clave"],
        API_ROWS,
        [Inches(1.2), Inches(1.5), Inches(2.7), Inches(1.1)],
    )

    add_heading(doc, "3. Organizacion del Proyecto")
    add_body(
        doc,
        "La estructura del repositorio separa claramente backend, frontend, base de datos, "
        "documentacion y automatizaciones. Eso facilita evolucionar Fase 0 y Fase 1 sin rehacer la base."
    )
    table2 = doc.add_table(rows=1, cols=3)
    fill_table(
        table2,
        ["Carpeta", "Rol", "Que contiene"],
        PROJECT_ROWS,
        [Inches(1.35), Inches(1.6), Inches(3.55)],
    )

    add_heading(doc, "4. Modulos Clave y Archivos Relevantes")
    table3 = doc.add_table(rows=1, cols=3)
    fill_table(
        table3,
        ["Modulo", "Archivos principales", "Funcion"],
        MODULE_ROWS,
        [Inches(1.25), Inches(2.55), Inches(2.7)],
    )

    add_heading(doc, "5. Vistas Principales del Frontend")
    add_body(
        doc,
        "El frontend tambien esta dividido por feature, lo que permite separar experiencia interna "
        "de La Vitamina, portal de cliente e integraciones operativas."
    )
    table4 = doc.add_table(rows=1, cols=3)
    fill_table(
        table4,
        ["Vista", "Ruta interna", "Archivos visibles"],
        FRONTEND_ROWS,
        [Inches(1.15), Inches(2.25), Inches(3.1)],
    )

    add_heading(doc, "6. Flujo Tecnico CRM -> Reservas -> Meta")
    add_bullets(
        doc,
        [
            "Un lead entra desde formulario, campaña, carga manual o futura integracion externa.",
            "El backend lo registra en CRM con organization_id como contenedor tecnico y client_id como alcance funcional.",
            "El equipo de La Vitamina califica, contacta y mueve el caso dentro del bloque CRM.",
            "Si el proceso requiere agenda, el lead o contacto se vincula a un formulario publico o gestion interna de reservas.",
            "La reserva genera trazabilidad operativa mediante estados, eventos y datos de asistencia.",
            "Las integraciones de Meta y reporting pueden leer ese recorrido para atribucion, conversion y seguimiento por cliente.",
        ],
    )

    add_heading(doc, "7. Lectura para Presentacion")
    add_body(
        doc,
        "La mejor forma de presentar este sistema hoy es como una base ya modularizada: "
        "el proyecto no parte de cero, sino que ya tiene separacion tecnica suficiente para "
        "ordenar operacion, CRM, reservas e integraciones alrededor del modelo de La Vitamina."
    )
    add_bullets(
        doc,
        [
            "Fase 0: base tecnica, autenticacion, roles, clientes, modulos y estructura de datos.",
            "Fase 1: CRM operativo, reservas conectadas, portal cliente e integracion con Meta/reporting.",
            "Siguiente maduracion: reforzar la trazabilidad entre lead, contacto, reserva, asistencia y conversion.",
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
