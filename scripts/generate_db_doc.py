from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\leno\Desktop\final\vitahub-platform")
SCHEMA_MD = ROOT / "database" / "documentation" / "schema.md"
DATABASE_MD = ROOT / "docs" / "DATABASE.md"
OUTPUT = ROOT / "docs" / "current" / "VITAHUB_Base_de_Datos_Estructura_y_Relaciones.docx"


PRIMARY = RGBColor(0x2E, 0x74, 0xB5)
DARK = RGBColor(0x1F, 0x4D, 0x78)
TEXT = RGBColor(0x22, 0x22, 0x22)
LIGHT = RGBColor(0xF2, 0xF4, 0xF7)


def set_font(run, name: str, size: int, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def style_paragraph(paragraph, before=0, after=6, line=1.1, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.0)
    run = p.add_run(text)
    set_font(run, "Calibri", size, bold=bold, color=color or TEXT)


def parse_schema_tables(text: str):
    tables = []
    pattern = re.compile(r"## Table: `([^`]+)`\n\n\*\*Purpose:\*\* (.+?)\n\n\| Column \| Type \| Nullable \| Default \| Description \|\n\|[-| ]+\|\n(.*?)(?:\n\*\*Indexes:\*\*|\n---)", re.S)
    for match in pattern.finditer(text):
        name, purpose, rows_blob = match.groups()
        rows = []
        for line in rows_blob.strip().splitlines():
            if not line.startswith("|"):
                continue
            parts = [part.strip() for part in line.strip().strip("|").split("|")]
            if len(parts) >= 5:
                rows.append(parts[:5])
        indexes_match = re.search(rf"## Table: `{re.escape(name)}`.*?\*\*Indexes:\*\*\n(.*?)(?:\n\*\*Foreign Keys:\*\*|\n---)", text, re.S)
        fk_match = re.search(rf"## Table: `{re.escape(name)}`.*?\*\*Foreign Keys:\*\*\n(.*?)(?:\n---|\Z)", text, re.S)
        indexes = []
        fks = []
        if indexes_match:
            indexes = [line.strip("- ").strip() for line in indexes_match.group(1).strip().splitlines() if line.strip()]
        if fk_match:
            fks = [line.strip("- ").strip() for line in fk_match.group(1).strip().splitlines() if line.strip()]
        tables.append({"name": name, "purpose": purpose.strip(), "rows": rows, "indexes": indexes, "fks": fks})
    return tables


def extract_relations(text: str):
    block_match = re.search(r"## Nucleo relacional vigente\n\n```text\n(.*?)```", text, re.S)
    if not block_match:
        return []
    return [line.strip() for line in block_match.group(1).splitlines() if line.strip()]


def add_title(doc: Document):
    p = doc.add_paragraph()
    style_paragraph(p, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run("VITAHUB - Base de Datos, Estructura y Relaciones")
    set_font(r, "Calibri", 22, bold=True, color=TEXT)

    p = doc.add_paragraph()
    style_paragraph(p, after=10, line=1.0)
    r = p.add_run("Documento tecnico generado desde la documentacion actual del proyecto.")
    set_font(r, "Calibri", 10, color=DARK)

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    widths = [Inches(1.8), Inches(4.7)]
    for row in meta.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
    entries = [
        ("Proyecto", "VITAHUB"),
        ("Fuente principal", "database/documentation/schema.md + docs/DATABASE.md"),
        ("Fecha", "28 de julio de 2026"),
    ]
    for i, (label, value) in enumerate(entries):
        shade_cell(meta.cell(i, 0), "E8EEF5")
        set_cell_text(meta.cell(i, 0), label, bold=True)
        set_cell_text(meta.cell(i, 1), value)


def add_section_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    style_paragraph(p, before=16 if level == 1 else 10, after=6, line=1.0)
    r = p.add_run(text)
    set_font(r, "Calibri", 16 if level == 1 else 13, bold=True, color=PRIMARY if level == 1 else DARK)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        style_paragraph(p, after=4, line=1.15)
        r = p.add_run(item)
        set_font(r, "Calibri", 11, color=TEXT)


def main():
    schema_text = SCHEMA_MD.read_text(encoding="utf-8", errors="ignore")
    database_text = DATABASE_MD.read_text(encoding="utf-8", errors="ignore")
    tables = parse_schema_tables(schema_text)
    relations = extract_relations(database_text)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    add_title(doc)

    add_section_heading(doc, "1. Modelo General")
    p = doc.add_paragraph()
    style_paragraph(p, after=8, line=1.15)
    r = p.add_run(
        "La base de datos de VITAHUB usa MySQL con TypeORM y mantiene `organization_id` como "
        "clave tecnica de aislamiento. En el modelo funcional actual, esa organizacion representa "
        "a La Vitamina como empresa operadora que administra multiples cuentas cliente dentro del "
        "mismo ecosistema CRM, reservas, Meta, produccion, reportes e integraciones."
    )
    set_font(r, "Calibri", 11, color=TEXT)
    add_bullets(doc, [
        "Motor actual: MySQL 8",
        "ORM: TypeORM",
        "Clave tecnica de aislamiento: organization_id",
        "Interpretacion funcional: La Vitamina opera y segmenta por cliente dentro de su propia estructura",
        "En Fase 1, el flujo reservas + CRM + Meta se apoya en client_id como scope funcional",
    ])

    add_section_heading(doc, "2. Relaciones Base")
    if relations:
        rel_table = doc.add_table(rows=1, cols=1)
        rel_table.style = "Table Grid"
        shade_cell(rel_table.cell(0, 0), "F2F4F7")
        set_cell_text(rel_table.cell(0, 0), "Nucleo relacional vigente", bold=True)
        for rel in relations:
            row = rel_table.add_row()
            set_cell_text(row.cells[0], rel, size=10)

    add_section_heading(doc, "3. Dominios Principales")
    domain_rows = [
        ("Tenancy y acceso", "organizations (contenedor tecnico de La Vitamina), users, refresh_tokens, parameter_definitions, parameter_values, audit_logs"),
        ("Clientes y CRM", "clients, leads, crm_contacts, crm_opportunities, crm_interactions"),
        ("Reservas", "reservation_forms, reservations, reservation_events, reservation_form_events, reservation_availability_blocks, reservation_coupons"),
        ("Produccion", "pieces, piece_versions, corrections, content_grids, content_items"),
        ("Presupuesto y gamificacion", "ud_budgets, ud_movements, xp_periods, xp_events"),
        ("Operacion y aprobaciones", "meetings, meeting_attendees, action_items, approval_requests, approval_decisions"),
        ("Comercial y facturacion", "services, quotes, quote_items, invoices, payments, contracts, briefs, onboarding"),
        ("Integraciones", "integrations, integration_accounts, sync_runs, documents, uploads, notifications"),
    ]
    dom_table = doc.add_table(rows=1, cols=2)
    dom_table.style = "Table Grid"
    headers = ["Dominio", "Tablas clave"]
    for i, header in enumerate(headers):
        shade_cell(dom_table.cell(0, i), "E8EEF5")
        set_cell_text(dom_table.cell(0, i), header, bold=True)
    for left, right in domain_rows:
        row = dom_table.add_row()
        set_cell_text(row.cells[0], left, bold=True)
        set_cell_text(row.cells[1], right)

    add_section_heading(doc, "4. Inventario De Tablas")
    p = doc.add_paragraph()
    style_paragraph(p, after=6, line=1.15)
    r = p.add_run("A continuacion se listan las tablas documentadas, con su proposito, columnas, indices y claves foraneas.")
    set_font(r, "Calibri", 11, color=TEXT)

    for idx, table in enumerate(tables):
        if idx and idx % 4 == 0:
            doc.add_section(WD_SECTION.NEW_PAGE)
        add_section_heading(doc, f"4.{idx + 1} {table['name']}", level=2)
        p = doc.add_paragraph()
        style_paragraph(p, after=6, line=1.1)
        r = p.add_run(table["purpose"])
        set_font(r, "Calibri", 11, color=TEXT)

        grid = doc.add_table(rows=1, cols=5)
        grid.style = "Table Grid"
        widths = [Inches(1.3), Inches(1.0), Inches(0.9), Inches(1.0), Inches(2.3)]
        labels = ["Columna", "Tipo", "Nullable", "Default", "Descripcion"]
        for j, label in enumerate(labels):
            cell = grid.cell(0, j)
            cell.width = widths[j]
            shade_cell(cell, "E8EEF5")
            set_cell_text(cell, label, bold=True)
        for row_vals in table["rows"]:
            row = grid.add_row()
            for j, val in enumerate(row_vals):
                row.cells[j].width = widths[j]
                set_cell_text(row.cells[j], val, size=9)

        if table["indexes"]:
            p = doc.add_paragraph()
            style_paragraph(p, before=4, after=2)
            r = p.add_run("Indices")
            set_font(r, "Calibri", 11, bold=True, color=DARK)
            add_bullets(doc, table["indexes"])

        if table["fks"]:
            p = doc.add_paragraph()
            style_paragraph(p, before=2, after=2)
            r = p.add_run("Claves foraneas")
            set_font(r, "Calibri", 11, bold=True, color=DARK)
            add_bullets(doc, table["fks"])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(str(OUTPUT))


if __name__ == "__main__":
    main()
